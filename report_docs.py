from report_tables import *
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def creation_report_docx():
    #открытие документа
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.bold = True
    style.font.size = Pt(12)
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Mm(10)
    section.right_margin = Mm(5)
    section.top_margin = Mm(2.5)
    section.bottom_margin = Mm(5)
    #paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #конструктор
    #заголовок
    head = document.add_heading('СПРАВКА')
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    head2 = document.add_heading('по отчёту СИР свыше 5 месяцев', level=2)
    head2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    #основной текст
    document.add_paragraph()
    document.add_paragraph(('В районе внуково на настоящий момент задолженность свыше 5 месяцев имеют {0} ФЛС на сумму {1} млн.рублей.').format(data['fl_317']['fls'], data['fl_317']['summa']))
    
    
    #сохранение документа
    itog_spr = 'Справка.docx'
    document.save(itog_spr)